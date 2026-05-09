#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import List, Tuple

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BODY_SIZE = 11
TITLE_COLOR = RGBColor(24, 57, 99)
PRIMARY_COLOR = RGBColor(31, 78, 121)
SECONDARY_COLOR = RGBColor(96, 96, 96)
NOTE_COLOR = RGBColor(161, 98, 7)
GRAY = RGBColor(110, 110, 110)

NOTE_SECTIONS = {
    'trước khi bắt đầu',
    'rủi ro hoặc giới hạn',
    'rủi ro',
    'giới hạn',
    'lưu ý',
    'ghi chú hỗ trợ',
}
STEP_SECTIONS = {
    'cách áp dụng',
    'các bước thực hiện',
    'hướng dẫn thực hiện',
    'thao tác',
}
BULLET_HIGHLIGHT_SECTIONS = {
    'tóm tắt ngắn',
    'phím tắt dùng nhanh',
    'trong tài liệu này',
}


def set_run_font(run, size=BODY_SIZE, bold=False, color=None):
    run.font.name = 'Arial'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size, bold=False, color=None):
    style.font.name = 'Arial'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    style.font.size = Pt(size)
    style.font.bold = bold
    if color is not None:
        style.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for edge, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def add_field(paragraph, instr, size=10):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), instr)
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    rfonts = OxmlElement('w:rFonts')
    for attr in ('ascii', 'hAnsi', 'eastAsia'):
        rfonts.set(qn(f'w:{attr}'), 'Arial')
    rpr.append(rfonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rpr.append(sz)
    run.append(rpr)
    t = OxmlElement('w:t')
    t.text = '1'
    run.append(t)
    fld.append(run)
    paragraph._p.append(fld)


def add_paragraph(doc, text='', style=None, align=None, before=0, after=6, line=1.15, color=None, bold=False, size=BODY_SIZE):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    return p


def add_bullets(doc, items, fill='F8FAFD'):
    table = doc.add_table(rows=len(items), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for idx, item in enumerate(items):
        cell = table.cell(idx, 0)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=70, start=140, bottom=70, end=140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        bullet = p.add_run('• ')
        set_run_font(bullet, size=11, bold=True, color=PRIMARY_COLOR)
        text_run = p.add_run(item)
        set_run_font(text_run, size=11)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_note_box(doc, title, items, fill='FFF7E8', title_color=NOTE_COLOR):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=title_color)
    for item in items:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        b = p.add_run('• ')
        set_run_font(b, size=11, bold=True, color=title_color)
        t = p.add_run(item)
        set_run_font(t, size=11)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_box(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, 'F5F7FA')
    set_cell_margins(cell, top=110, start=140, bottom=110, end=140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for i, line in enumerate(lines):
        r = p.add_run(line)
        set_run_font(r, size=10, color=RGBColor(44, 62, 80))
        if i < len(lines) - 1:
            r.add_break()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_step_boxes(doc, steps):
    for idx, step in enumerate(steps, 1):
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        left_width = Cm(2.35)
        right_width = Cm(12.65)
        table.columns[0].width = left_width
        table.columns[1].width = right_width
        left = table.cell(0, 0)
        right = table.cell(0, 1)
        left.width = left_width
        right.width = right_width
        set_cell_shading(left, '1F4E79')
        set_cell_shading(right, 'F6F8FB')
        set_cell_margins(left, top=90, start=90, bottom=90, end=90)
        set_cell_margins(right, top=90, start=150, bottom=90, end=150)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p_left = left.paragraphs[0]
        p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_left.paragraph_format.space_after = Pt(0)
        label = p_left.add_run(f'Bước {idx}')
        set_run_font(label, size=10, bold=True, color=RGBColor(255, 255, 255))
        p_right = right.paragraphs[0]
        p_right.paragraph_format.space_after = Pt(0)
        p_right.paragraph_format.line_spacing = 1.1
        r_right = p_right.add_run(step)
        set_run_font(r_right, size=11)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_section_title(doc, index, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    num = p.add_run(f'{index}. ')
    set_run_font(num, size=13, bold=True, color=PRIMARY_COLOR)
    txt = p.add_run(title)
    set_run_font(txt, size=13, bold=True, color=PRIMARY_COLOR)


def slug_to_title(path: Path) -> str:
    name = path.stem
    name = re.sub(r'^\d{4}-\d{2}-\d{2}[_-]?', '', name)
    name = name.replace('_', ' ').replace('-', ' ')
    return re.sub(r'\s+', ' ', name).strip().title()


def parse_markdown(path: Path) -> Tuple[str | None, List[Tuple[str, list]]]:
    title = None
    sections: List[Tuple[str, list]] = []
    current_title = None
    current_items = []
    in_code = False
    code_lines: List[str] = []

    for raw in path.read_text(encoding='utf-8').splitlines():
        line = raw.rstrip('\n')
        stripped = line.strip()
        if line.startswith('```'):
            if in_code:
                current_items.append(('code', code_lines[:]))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if stripped.startswith('# '):
            title = stripped[2:].strip()
            continue
        if stripped.startswith('## '):
            if current_title is not None:
                sections.append((current_title, current_items))
            current_title = stripped[3:].strip()
            current_items = []
            continue
        if not stripped:
            continue
        if stripped.startswith('- '):
            current_items.append(('bullet', stripped[2:].strip()))
        else:
            current_items.append(('text', stripped))
    if current_title is not None:
        sections.append((current_title, current_items))
    return title, sections


def collect_steps(items: list) -> List[str]:
    steps: List[str] = []
    for kind, value in items:
        if kind == 'bullet':
            step_text = re.sub(r'^Bước\s*\d+\s*:\s*', '', value, flags=re.IGNORECASE).strip()
            steps.append(step_text)
    return steps


def render(doc: Document, title: str, sections: List[Tuple[str, list]], header_prefix: str):
    add_paragraph(doc, 'HƯỚNG DẪN SỬ DỤNG', align=WD_ALIGN_PARAGRAPH.CENTER, after=4, line=1.0, color=GRAY, bold=True, size=10)
    add_paragraph(doc, title, style='Title', align=WD_ALIGN_PARAGRAPH.CENTER, after=6, line=1.0, color=TITLE_COLOR, bold=True, size=18)
    add_paragraph(doc, 'Tài liệu này được trình bày theo phong cách hướng dẫn sử dụng, ưu tiên thao tác rõ ràng và dễ quét.', align=WD_ALIGN_PARAGRAPH.CENTER, after=12, color=SECONDARY_COLOR)
    add_paragraph(doc, '_' * 62, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, line=1.0, color=RGBColor(210, 214, 220), size=8)

    for index, (section_title, items) in enumerate(sections, 1):
        normalized = section_title.lower().strip()
        add_section_title(doc, index, section_title)
        if normalized in STEP_SECTIONS:
            intro = [v for k, v in items if k == 'text']
            for paragraph in intro[:1]:
                add_paragraph(doc, paragraph, after=8)
            steps = collect_steps(items)
            if steps:
                add_step_boxes(doc, steps)
            remainder_codes = [v for k, v in items if k == 'code']
            for code in remainder_codes:
                add_code_box(doc, code)
            other_bullets = [v for k, v in items if k == 'bullet' and not re.match(r'^Bước\s*\d+\s*:', v, flags=re.IGNORECASE)]
            if other_bullets:
                add_bullets(doc, other_bullets)
            continue
        if normalized in NOTE_SECTIONS:
            bullet_items = [v for k, v in items if k == 'bullet']
            text_items = [v for k, v in items if k == 'text']
            title_box = 'Lưu ý' if 'rủi ro' in normalized or 'giới hạn' in normalized else section_title
            content = bullet_items + text_items
            if content:
                fill = 'EEF6FF' if normalized == 'trước khi bắt đầu' else 'FFF7E8'
                color = PRIMARY_COLOR if normalized == 'trước khi bắt đầu' else NOTE_COLOR
                add_note_box(doc, title_box, content, fill=fill, title_color=color)
            for code in [v for k, v in items if k == 'code']:
                add_code_box(doc, code)
            continue
        if normalized in BULLET_HIGHLIGHT_SECTIONS:
            bullets = [v for k, v in items if k == 'bullet']
            texts = [v for k, v in items if k == 'text']
            if bullets:
                add_bullets(doc, bullets)
            for text in texts:
                add_paragraph(doc, text)
            for code in [v for k, v in items if k == 'code']:
                add_code_box(doc, code)
            continue
        for kind, value in items:
            if kind == 'text':
                add_paragraph(doc, value)
            elif kind == 'bullet':
                add_bullets(doc, [value])
            elif kind == 'code':
                add_code_box(doc, value)


def build_document(input_path: Path, output_path: Path, title: str, header_prefix: str, margin_cm: float, body_size: int):
    global BODY_SIZE
    BODY_SIZE = body_size
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)

    for style_name, size, bold, color in [
        ('Normal', body_size, False, None),
        ('Title', 18, True, TITLE_COLOR),
        ('Heading 1', 13, True, PRIMARY_COLOR),
        ('Heading 2', 12, True, PRIMARY_COLOR),
        ('List Bullet', body_size, False, None),
    ]:
        set_style_font(doc.styles[style_name], size, bold, color)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(6)
    header_run = header_p.add_run(f'{header_prefix} {title}'.strip())
    set_run_font(header_run, size=10, bold=False, color=GRAY)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(4)
    footer_p.paragraph_format.space_after = Pt(0)
    footer_p.paragraph_format.line_spacing = 1.0
    run = footer_p.add_run('Trang ')
    set_run_font(run, size=10, bold=False, color=GRAY)
    add_field(footer_p, 'PAGE', size=10)
    run = footer_p.add_run('/')
    set_run_font(run, size=10, bold=False, color=GRAY)
    add_field(footer_p, 'NUMPAGES', size=10)

    parsed_title, sections = parse_markdown(input_path)
    render(doc, title or parsed_title or slug_to_title(input_path), sections, header_prefix)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description='Generate DOCX HDSD from Markdown note.')
    parser.add_argument('input', help='Path to source markdown file')
    parser.add_argument('--output', help='Output .docx path')
    parser.add_argument('--title', help='Override title shown in document')
    parser.add_argument('--header-prefix', default='HDSD', help='Prefix shown in header')
    parser.add_argument('--margin-cm', type=float, default=2.0, help='Page margins in cm')
    parser.add_argument('--body-font-size', type=int, default=11, help='Body font size')
    args = parser.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        raise SystemExit(f'Input file not found: {input_path}')
    output_path = Path(args.output).expanduser().resolve() if args.output else input_path.with_suffix('.docx')
    title = args.title
    if title is None:
        parsed_title, _ = parse_markdown(input_path)
        title = parsed_title or slug_to_title(input_path)
    try:
        build_document(input_path, output_path, title, args.header_prefix, args.margin_cm, args.body_font_size)
    except PermissionError:
        fallback = output_path.with_name(output_path.stem + '_new' + output_path.suffix)
        build_document(input_path, fallback, title, args.header_prefix, args.margin_cm, args.body_font_size)
        output_path = fallback
        print(f'[WARN] Output file was locked. Saved to: {output_path}')
        return
    print(output_path)


if __name__ == '__main__':
    main()
